from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "output" / "doc" / "基于Go与Uni-第5第6章补充版.docx"
OUT = ROOT / "output" / "doc" / "基于Go与Uni-模块展开修正版.docx"
BACKUP = ROOT / "output" / "doc" / "基于Go与Uni-模块展开修正版前备份.docx"
ASSETS = ROOT / "output" / "doc" / "assets"


def set_run_font(run, size=12):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
        "宋体",
    )
    run.font.size = Pt(size)


def set_paragraph_font(paragraph, size=12):
    for run in paragraph.runs:
        set_run_font(run, size=size)


def body_paragraph(paragraph):
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.5
    set_paragraph_font(paragraph)


def caption_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    set_paragraph_font(paragraph)


class Inserter:
    def __init__(self, doc, target):
        self.doc = doc
        self.target = target

    def move_before_target(self, obj):
        element = obj._tbl if hasattr(obj, "_tbl") else obj._p
        self.target.addprevious(element)
        return obj

    def p(self, text="", style=None, indent=True):
        paragraph = self.doc.add_paragraph(text)
        if style:
            paragraph.style = style
        if style is None and indent and text:
            body_paragraph(paragraph)
        elif text:
            set_paragraph_font(paragraph)
        return self.move_before_target(paragraph)

    def heading(self, text, level):
        paragraph = self.doc.add_paragraph(text, style=f"Heading {level}")
        set_paragraph_font(paragraph, size=16 if level == 1 else 14 if level == 2 else 12)
        return self.move_before_target(paragraph)

    def page_break(self):
        paragraph = self.doc.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        return self.move_before_target(paragraph)

    def table_caption(self, text):
        paragraph = self.doc.add_paragraph(text)
        caption_paragraph(paragraph)
        return self.move_before_target(paragraph)

    def image(self, filename, caption, width=5.7):
        path = ASSETS / filename
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run().add_picture(str(path), width=Inches(width))
        self.move_before_target(paragraph)
        cap = self.doc.add_paragraph(caption)
        caption_paragraph(cap)
        return self.move_before_target(cap)

    def table(self, rows):
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            for j, value in enumerate(row):
                cell = table.cell(i, j)
                cell.text = value
                for p in cell.paragraphs:
                    set_paragraph_font(p, size=10.5)
                    if i == 0:
                        for run in p.runs:
                            run.bold = True
        return self.move_before_target(table)


def find_paragraph(doc, prefix):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def remove_range(doc, start_prefix, end_prefix):
    start = find_paragraph(doc, start_prefix)._p
    end = find_paragraph(doc, end_prefix)._p
    body = doc._body._element
    children = list(body)
    start_idx = children.index(start)
    end_idx = children.index(end)
    for element in children[start_idx:end_idx]:
        body.remove(element)
    return end


def add_chapter_4(ins):
    ins.heading("第四章  系统设计", 1)
    ins.heading("4.1 系统总体设计", 2)
    ins.p("酒店管理系统采用前后端分离与分层服务相结合的总体架构。系统面向两类主要使用场景：一类是住客通过 UniApp 小程序端完成酒店浏览、房型筛选、在线预订、订单查看和个人资料维护；另一类是酒店工作人员通过 React 管理端完成订单确认、前台入住退房、房间维护、库存价格配置、工单处理、公告活动发布和系统基础资料维护。后端使用 Go 语言和 Gin 框架提供统一接口，业务规则集中在 Service 层，数据访问由 Repository 与 GORM 负责。")
    ins.p("从系统边界看，小程序端和管理端并不直接访问数据库，而是统一通过 HTTP 接口与后端交互。后端在公共路由中提供房间查询、活动横幅、公告、酒店设置和推荐结果等对外展示能力，在鉴权路由中提供用户资料、预订订单和行为上报等住客能力，在管理员路由中提供订单、房间、房型、库存、定价、工单、用户、日志和系统设置等后台能力。该划分使用户端展示、后台运营和底层数据之间保持清晰边界。")
    ins.image("fig4-1-softeng-architecture.png", "图 4-1 系统总体架构图", 5.9)
    ins.p("如图 4-1 所示，系统自上而下可划分为表示层、接口接入层、业务服务层、数据访问层和基础支撑层。表示层包括 UniApp 小程序端与 React 管理端；接口接入层负责路由分组、JWT 鉴权、跨域处理和日志上报；业务服务层承载用户、房间、订单、库存、定价、推荐、工单、内容和系统设置等核心规则；数据访问层封装用户表、房间表、订单表、库存表、定价表、工单表、公告表和日志表等数据操作；基础支撑层则包含 MySQL/SQLite、对象存储、Swagger 接口文档和多层时间轮任务。")

    ins.heading("4.2 系统功能模块设计", 2)
    ins.p("参考系统实际页面、接口和服务代码，本文将功能模块按使用对象和业务职责进行划分，而不是只概括为预订模块和后台管理模块。系统功能模块整体包括住客端功能、前台接待与订单账务功能、房态运维与库存价格功能、房务工单功能、内容运营功能以及系统支撑功能。系统总体功能模块如图 4-2 所示。")
    ins.image("fig4-2-softeng-modules.png", "图 4-2 系统总体功能模块图", 5.9)

    ins.heading("4.2.1 住客端功能模块", 3)
    ins.p("住客端功能模块主要部署在 UniApp 小程序端，对应 pages.json 中的首页、酒店、服务、我的、登录注册、房间详情、确认订单、订单详情、个人资料、账号安全、余额、积分、优惠券、关于、隐私政策和用户协议等页面。该端的设计目标是让住客在移动端完成从了解酒店到提交订单再到查看入住信息的完整闭环。")
    ins.table_caption("表 4-1 住客端功能模块设计")
    ins.table([
        ["模块名称", "对应页面或接口", "主要职责"],
        ["用户认证模块", "login、register、/api/auth/login、/api/auth/register", "完成住客注册、登录、令牌保存和受保护接口访问。"],
        ["首页与酒店展示模块", "index、service、/api/settings/public、/api/banners/active、/api/notices/active", "展示酒店名称、地址、横幅活动、公告、服务信息和快捷入口。"],
        ["房型浏览模块", "hotel、room-detail、/api/rooms、/api/rooms/available、/api/room-categories", "展示可售房间、房型分类、图片、设施、床型、面积和基础价格。"],
        ["推荐行为模块", "/api/recommendations/rooms、/api/recommendations/behavior", "记录浏览、详情点击和预订意图，生成推荐房型和推荐理由。"],
        ["在线预订模块", "booking、pay-result、/api/bookings", "填写入住人信息、校验日期库存、计算价格并创建订单。"],
        ["订单管理模块", "orders、order-detail、/api/bookings/my、/api/bookings/:id/cancel", "查看订单列表、订单详情、状态流转和可取消订单。"],
        ["会员资料模块", "my、profile、security、wallet、points、coupons、/api/users/profile", "维护个人资料、修改密码，展示会员等级、余额、积分和优惠券入口。"],
    ])
    ins.p("需要说明的是，优惠券页面在当前版本中主要作为会员资产入口保留，后端没有将优惠券接入真实订单抵扣流程，因此本文只将其描述为会员中心的扩展入口，不把它写成已经完成的优惠券核销或结算模块。")

    ins.heading("4.2.2 前台接待与订单账务模块", 3)
    ins.p("前台接待与订单账务模块对应管理端的订单管理、客人办理入住和客人办理退房页面。该模块连接住客端预订结果和酒店线下接待流程，工作人员可以查看全部订单、根据客人姓名或手机号检索预订、确认订单、办理入住、办理退房和处理取消订单。订单状态变化会同步影响房间状态、库存数量、清洁状态和会员积分，是系统业务闭环中最关键的管理端模块。")
    ins.table_caption("表 4-2 前台接待与订单账务模块设计")
    ins.table([
        ["模块名称", "对应页面或接口", "主要职责"],
        ["订单管理", "order-manage、/api/admin/bookings", "分页查看订单、确认订单、检索订单和跟踪订单状态。"],
        ["前台入住", "hotel-front/check-in、/api/admin/bookings/search、/api/admin/bookings/:id/checkin", "按客人信息检索预订，核对订单后将房间置为入住状态。"],
        ["前台退房", "hotel-front/check-out、/api/admin/bookings/:id/checkout", "办理退房，恢复房态并将清洁状态置为待清洁。"],
        ["订单取消", "/api/bookings/:id/cancel", "住客取消订单后释放对应库存，后台保持订单状态一致。"],
        ["账务与会员联动", "BookingService、UserService", "退房后累计消费金额，更新积分和会员等级。"],
    ])

    ins.heading("4.2.3 房态运维与库存价格模块", 3)
    ins.p("房态运维与库存价格模块面向酒店内部运营人员，包含房间列表、批量创建房间、房型分类、楼层可视化、设施摆放、房态库存矩阵和动态定价规则等功能。系统将房间基础资料、房型主数据、日期库存和价格规则拆分保存，使用户端房型展示、后台库存维护和订单价格计算能够共享同一套数据来源。")
    ins.table_caption("表 4-3 房态运维与库存价格模块设计")
    ins.table([
        ["模块名称", "对应页面或接口", "主要职责"],
        ["房间管理", "room-manage/list、/api/rooms", "维护房号、楼层、房型、价格、面积、床型、房态和清洁状态。"],
        ["房型分类", "room-manage/category、/api/admin/room-categories", "维护房型名称、描述、设施和图片，并同步相关房间与库存字段。"],
        ["房间可视化", "room-visualization、/api/admin/facilities", "在楼层平面中拖拽房间和公共设施，提高房态查看直观性。"],
        ["库存矩阵", "inventory-manage、/api/admin/inventory/grid", "按房型和日期查看总房量、已订房量、剩余房量和基础价格。"],
        ["动态定价", "pricing-manage、/api/admin/pricing/rules", "按日期区间、规则类型和优先级配置加价、减价或百分比调整。"],
    ])

    ins.heading("4.2.4 房务工单、内容运营与系统支撑模块", 3)
    ins.p("除预订和房态外，系统还实现了维修清洁工单、横幅活动、公告提醒、文件上传、后台用户、操作日志和系统设置等支撑模块。这些模块虽然不一定直接参与订单创建，但会影响房间可售状态、住客端展示内容、后台权限管理和系统运维能力。")
    ins.table_caption("表 4-4 房务工单、内容运营与系统支撑模块设计")
    ins.table([
        ["模块名称", "对应页面或接口", "主要职责"],
        ["维修工单", "work-order/repair、/api/admin/work-orders/repair", "登记维修任务，维修期间将房间状态置为 maintenance。"],
        ["清洁任务", "work-order/cleaning、/api/admin/work-orders/cleaning", "分配和完成清洁任务，完成后恢复房间清洁状态。"],
        ["营销活动", "activity-manage、/api/admin/banners", "维护首页横幅、图片、排序、启停状态和展示时间窗。"],
        ["公告提醒", "notice-manage、/api/admin/notices", "维护住客端公告内容，并通过时间窗控制展示状态。"],
        ["系统设置", "system-setting、/api/admin/settings/save", "维护酒店资料、客服电话、入住规则、地图坐标和封面图。"],
        ["后台用户", "user-manage、/api/admin/users", "维护管理员和普通用户资料、状态、角色和默认密码。"],
        ["日志与上传", "log-manage、/api/logs/report、/api/upload/image", "记录前端日志、后台操作线索，并为图片资源提供上传入口。"],
    ])

    ins.heading("4.3 核心业务流程设计", 2)
    ins.p("预订与入住流程是系统中跨模块最多的业务链路。住客先在小程序端选择日期并浏览房型，进入详情页后提交预订信息；后端在创建订单时校验用户身份、房间状态和日期库存，并计算订单金额；管理员确认订单后，前台工作人员办理入住；退房后系统更新房态、清洁状态、库存和会员积分。预订与入住业务流程如图 4-3 所示。")
    ins.image("fig4-3-softeng-booking-flow.png", "图 4-3 预订与入住业务流程图", 6.0)
    ins.p("库存与价格联动是房型展示和订单创建的基础。系统按房型和日期保存库存记录，在展示可售房间和创建订单时读取对应日期区间的库存；动态定价规则按房型、日期区间和优先级参与每日价格计算；订单创建、取消和退房等操作会触发库存或房态回写。该联动关系如图 4-4 所示。")
    ins.image("fig4-4-softeng-inventory-pricing.png", "图 4-4 库存与动态定价联动流程图", 6.0)

    ins.heading("4.4 本章小结", 2)
    ins.p("本章按照参考论文的模块化表达方式，对酒店管理系统的总体架构、功能模块和核心业务流程进行了重新梳理。与原先只概括住客服务和运营管理两个大类不同，本章进一步列出了用户认证、首页展示、房型浏览、推荐、在线预订、订单管理、前台接待、房间管理、房型分类、房态可视化、库存矩阵、动态定价、维修清洁、活动公告、系统设置、后台用户、日志上传等实际模块，为后续数据库设计、详细实现和测试章节提供清晰依据。")


def add_chapter_5(ins):
    ins.page_break()
    ins.heading("第五章  数据库设计", 1)
    ins.heading("5.1 数据库概念结构设计", 2)
    ins.p("数据库设计需要与第四章列出的功能模块对应。根据后端 models 目录和自动迁移逻辑，系统围绕 users、rooms、room_categories、bookings、room_inventories、pricing_rules、facilities、maintenance、housekeeping、banners、notices、hotel_settings、logs、user_behaviors 和 hotels 等实体组织数据。各实体并非孤立存在，而是共同支撑住客预订、房态维护、库存价格、房务工单、内容展示和后台管理等业务场景。")
    ins.image("er-redesign-drawio.png", "图 5-1 系统数据库E-R图", 6.0)
    ins.p("如图 5-1 所示，用户实体与订单实体之间为一对多关系；房间实体与订单实体之间为一对多关系；房型分类实体与房间、库存、定价规则之间通过房型名称保持关联；库存实体按房型和日期记录可售能力；价格规则实体作用于库存基础价；维修与清洁实体围绕房间状态变化展开；横幅、公告和系统设置实体服务于小程序端内容展示；日志和行为记录实体分别服务于系统运维与推荐计算。")

    ins.heading("5.2 数据库逻辑结构设计", 2)
    ins.table_caption("表 5-1 核心实体与功能模块对应关系")
    ins.table([
        ["实体或数据表", "对应功能模块", "主要用途"],
        ["users", "用户认证、会员资料、后台用户", "保存账号、角色、状态、会员等级、积分、余额和累计消费。"],
        ["rooms", "房间管理、房态可视化、订单入住", "保存房号、楼层、房型、价格、房态、清洁状态和可视化坐标。"],
        ["room_categories", "房型分类、房型展示", "保存房型描述、设施、图片和使用数量统计依据。"],
        ["bookings", "在线预订、订单管理、前台接待", "保存入住人、日期、金额、订单状态、支付状态和房间关联。"],
        ["room_inventories", "库存矩阵、可售查询", "按房型和日期保存总房量、已订房量、剩余房量和基础价格。"],
        ["pricing_rules", "动态定价", "保存加价、减价、百分比调整、日期范围和优先级。"],
        ["maintenance、housekeeping", "维修清洁工单", "记录维修任务、清洁任务、处理状态、负责人和完成时间。"],
        ["banners、notices", "营销活动、公告提醒", "保存横幅图片、公告内容、排序、状态和展示时间窗。"],
        ["hotel_settings、hotels", "系统设置、酒店资料展示", "保存客服电话、地址、入住规则、地图坐标和酒店基础资料。"],
        ["logs、user_behaviors", "日志审计、推荐行为", "保存前端日志、后台操作线索和用户浏览、点击、预订意图。"],
    ])

    ins.p("逻辑结构设计中，系统没有把所有字段都压在一张大表中，而是按照业务边界拆分实体。例如，rooms 表记录单个房间的基础信息和状态，room_categories 表保存房型主数据，room_inventories 表按日期保存库存，pricing_rules 表表达价格调整策略。这样可以避免房间信息、房型描述、库存日期和价格规则相互混杂，也方便后台分别维护。")
    ins.table_caption("表 5-2 主要数据表逻辑结构")
    ins.table([
        ["表名", "关键字段", "设计说明"],
        ["users", "id、username、phone、role、status、level、points、balance、token_version", "支撑登录鉴权、会员资料、后台用户管理和旧令牌失效。"],
        ["rooms", "id、room_number、floor、room_type、price、status、clean_status、left、top", "支撑房源展示、房态流转和楼层可视化。"],
        ["room_categories", "id、name、description、facilities、images、status", "支撑房型统一展示和房型主数据维护。"],
        ["bookings", "id、user_id、room_id、guest_name、check_in_date、check_out_date、total_price、status", "支撑预订、取消、确认、入住和退房全流程。"],
        ["room_inventories", "id、room_type、date、total_count、booked_count、available_count、base_price", "支撑按日期查询库存和订单库存扣减。"],
        ["pricing_rules", "id、room_type、start_date、end_date、rule_type、adjustment、priority", "支撑节假日、周末和活动期间动态价格调整。"],
        ["maintenance", "id、room_id、repair_type、description、status、reporter_id、worker_id", "支撑维修登记、维修完成和房间维修状态联动。"],
        ["housekeeping", "id、room_id、cleaning_type、status、staff_id、completed_at", "支撑清洁任务分配、完成和清洁状态恢复。"],
        ["banners、notices", "title、image_url、content、sort_order、status、start_time、end_time", "支撑活动横幅和公告按时间窗发布。"],
        ["logs、user_behaviors", "level、message、source、behavior_type、room_id、user_id", "支撑日志查看、异常排查和推荐行为统计。"],
    ])

    ins.heading("5.3 业务状态与一致性设计", 2)
    ins.p("酒店业务中大量操作都体现为状态变化，因此状态字段设计是数据库设计的重要部分。订单状态、房间状态、清洁状态、支付状态、维修状态和公告活动状态共同决定系统页面展示和后端规则判断。如果这些状态口径不一致，就会出现用户端仍显示可订、后台却已经入住，或退房后房间未进入待清洁状态等问题。")
    ins.table_caption("表 5-3 主要状态字段设计")
    ins.table([
        ["状态字段", "所在实体", "主要取值或含义"],
        ["role、status", "users", "区分管理员、员工、普通住客以及账号启用、禁用状态。"],
        ["status", "rooms", "表达 available、occupied、maintenance 等房态。"],
        ["clean_status", "rooms", "表达 clean、dirty 等清洁状态，退房后通常置为 dirty。"],
        ["status", "bookings", "表达 pending、confirmed、checkin、checkout、cancelled 等订单阶段。"],
        ["payment_status", "bookings", "表达待支付、已支付或支付相关状态，服务于账务展示。"],
        ["status", "maintenance、housekeeping", "表达待处理、处理中、已完成等工单流转阶段。"],
        ["status、start_time、end_time", "banners、notices", "共同决定横幅和公告当前是否可展示。"],
    ])
    ins.p("为保证数据一致性，系统在 Service 层增加了多处业务校验。例如，房型分类更新时同步处理房间、库存和定价规则中使用的房型名称；删除房型前检查其是否仍被房间、库存或价格规则引用；创建订单前校验房间可用性和日期库存；订单取消时释放库存；入住后更新房间状态；退房后恢复房态并生成待清洁状态；维修工单创建和完成时也会同步影响房间状态。该设计使数据库结构与系统业务流程保持一致。")

    ins.heading("5.4 本章小结", 2)
    ins.p("本章从概念结构、逻辑结构和业务状态三个层面对数据库设计进行了说明。相比原稿只选取少量核心表进行概括，本章进一步列出了房型分类、库存矩阵、动态定价、维修清洁、横幅公告、系统设置、日志和用户行为等模块所需的数据表，并说明了这些表与第四章功能模块之间的对应关系，为第六章逐模块实现说明提供数据基础。")


def add_chapter_6(ins):
    ins.page_break()
    ins.heading("第六章  系统详细设计与实现", 1)
    ins.heading("6.1 小程序住客端详细设计与实现", 2)
    ins.heading("6.1.1 用户注册登录与账号安全模块", 3)
    ins.p("用户注册登录模块是住客端进入受保护业务的入口。住客在注册页面填写用户名、手机号、邮箱和密码后，前端调用 /api/auth/register 接口创建普通用户账号；登录页面调用 /api/auth/login 接口完成身份校验，登录成功后将 token 和用户资料保存到本地存储。后续访问个人中心、确认订单、我的订单和账号安全等页面时，请求拦截器会自动在 Authorization 头中携带 Bearer Token。")
    ins.p("账号安全模块由个人资料和修改密码两个页面组成。个人资料页面通过 /api/users/profile 获取和更新真实姓名、手机号和头像；修改密码页面通过 /api/users/password 提交旧密码和新密码。后端在 JWT 中写入 TokenVersion 字段，用户修改密码或账号状态变化后，旧令牌会因为版本不一致而失效，从而降低长期登录态带来的风险。")

    ins.heading("6.1.2 首页与酒店信息展示模块", 3)
    ins.p("首页模块主要负责把酒店基础信息、活动横幅、入住日期选择和快捷入口集中展示给住客。页面加载时读取酒店设置、活动横幅和公告信息；用户选择入住日期与离店日期后，点击立即预订即可跳转到房型列表页。首页同时提供我的订单、联系前台、地图导航和酒店服务等入口，使住客不必在多个页面之间反复查找常用功能。")
    ins.image("fig6-9-mini-home-room.png", "图 6-1 小程序端首页与房型列表界面", 5.8)

    ins.heading("6.1.3 房型浏览、详情与推荐模块", 3)
    ins.p("房型浏览模块围绕可售房间列表、房型分类和房间详情展开。房型列表页根据日期条件调用 /api/rooms/available 或房型查询接口，展示房型名称、面积、床型、可住人数、设施和起订价格。房间详情页进一步展示房间图片、预订说明和推荐房型，并在用户浏览详情、点击推荐和产生预订意图时上报用户行为。")
    ins.p("推荐模块并未实现复杂训练模型，而是根据用户行为、历史订单、房型价格、面积、可住人数、折扣率和热度等信息进行轻量化打分。系统通过 /api/recommendations/behavior 记录 view_room、view_detail、click_recommendation 和 book_intent 等行为，通过 /api/recommendations/rooms 返回推荐房型和推荐理由。")
    ins.image("fig6-10-mini-room-booking.png", "图 6-2 小程序端房间详情与确认订单界面", 5.8)

    ins.heading("6.1.4 在线预订与订单提交模块", 3)
    ins.p("在线预订模块是住客端业务闭环的核心。确认订单页会复用房间详情、入住日期、离店日期和用户资料，要求住客填写入住人姓名、手机号、身份证号和特殊需求。提交订单后，后端 BookingService 会校验日期范围、房间可用性和库存数量，再按每日库存基础价与动态定价规则累加计算订单金额。订单创建成功后，系统减少对应日期区间的库存数量，并返回订单信息。")

    ins.heading("6.1.5 订单查询、会员中心与酒店服务模块", 3)
    ins.p("订单查询模块包括订单列表和订单详情两个页面。订单列表按照全部、待确认、待入住、入住中、已完成和已取消等状态筛选订单；订单详情页展示入住人、房间、入住离店日期、订单金额、支付状态和酒店联系方式，并提供取消订单或查看房型入口。该模块使住客能够及时了解订单从创建到入住、退房或取消的状态变化。")
    ins.image("fig6-11-mini-order-flow.png", "图 6-3 小程序端订单列表与订单详情界面", 5.8)
    ins.p("会员中心和酒店服务模块补足了住客在预订之外的持续使用场景。个人中心展示会员等级、余额、积分和优惠券入口，并提供我的订单、个人资料、账号安全、酒店服务和关于我们等菜单；酒店服务页展示客服电话、酒店位置、WiFi 信息、行李寄送、入住退房时间、公告和设施开放时间。小程序端个人中心与酒店服务界面如图 6-4 所示。")
    ins.image("fig6-12-mini-profile-service.png", "图 6-4 小程序端个人中心与酒店服务界面", 5.8)

    ins.heading("6.2 管理端运营后台详细设计与实现", 2)
    ins.heading("6.2.1 登录鉴权与运营控制台模块", 3)
    ins.p("管理端登录模块与住客端共用后端认证接口，但页面面向酒店工作人员。管理员可通过账号密码或手机号完成登录，登录成功后进入运营控制台。控制台集中展示入住率、在住间数、待处理工单、订单和房态快捷入口，便于值班人员快速掌握当日运营情况。")
    ins.image("dashboard-refreshed-body.png", "图 6-5 管理端运营控制台界面", 5.8)
    ins.image("admin-login.png", "图 6-6 管理端登录界面", 5.8)

    ins.heading("6.2.2 订单管理与前台接待模块", 3)
    ins.p("订单管理模块用于后台查看和处理所有预订记录。工作人员可以通过订单列表了解订单号、住客信息、房间、入住日期、订单金额和订单状态，并执行确认订单等操作。前台接待模块在此基础上提供分步式办理入住与退房流程：办理入住时根据客人姓名或手机号检索订单，核对后将订单状态更新为 checkin，同时将房间状态置为 occupied；办理退房时将订单状态更新为 checkout，并把房间恢复为 available、清洁状态置为 dirty。")
    ins.image("order-manage.png", "图 6-7 订单管理界面", 5.8)
    ins.image("front-checkin.png", "图 6-8 前台办理入住界面", 5.8)

    ins.heading("6.2.3 房间列表、房型分类与房间可视化模块", 3)
    ins.p("房间管理模块维护酒店实际房源。房间列表支持新增、编辑、删除、批量创建和按房型、房态、清洁状态筛选；房型分类模块维护标准间、豪华套房、总统套房等房型主数据，并统一管理房型描述、图片和设施标签。由于房型名称会被房间、库存和定价规则共同引用，后端在修改和删除房型时会执行同步更新或引用检查。")
    ins.p("房间可视化模块将房间和公共设施抽象为带坐标与尺寸的平面对象。管理人员可以在楼层平面中拖拽房间、调整位置，也可以添加电梯、走廊、前台和餐饮区等设施，使房态管理更接近真实酒店场景。")
    ins.image("room-visualization.png", "图 6-9 房间可视化管理界面", 5.8)

    ins.heading("6.2.4 库存矩阵与动态定价模块", 3)
    ins.p("库存矩阵模块以日期为横轴、房型为纵轴，展示未来一段时间内不同房型的总房量、已订房量、剩余房量和基础价格。管理员可通过该界面快速识别高峰日期、紧缺房型和价格波动情况。库存服务提供初始化库存、增减库存和按日期区间查询库存网格等接口，并在订单创建、取消和退房时与 BookingService 联动。")
    ins.image("inventory-grid.png", "图 6-10 房态库存矩阵界面", 5.8)
    ins.p("动态定价模块以 PricingRule 为核心对象，支持按照房型、日期区间、规则类型、调整金额或百分比和优先级维护价格规则。系统计算订单价格时逐日读取库存基础价，再叠加符合条件的价格规则，使周末、节假日和活动期间的价格变化能够反映到订单金额中。")
    ins.image("pricing-rules.png", "图 6-11 动态定价规则管理界面", 5.8)

    ins.heading("6.2.5 维修清洁工单模块", 3)
    ins.p("维修清洁工单模块用于处理房务状态变化。创建维修工单后，系统会把对应房间状态置为 maintenance，防止该房间继续被预订；维修完成后，房间状态恢复为 available，并将清洁状态置为 dirty，提示后续清洁。清洁任务按照 pending、in_progress、completed 等状态流转，完成后将房间清洁状态更新为 clean。")
    ins.image("repair-orders.png", "图 6-12 工单处理管理界面", 5.8)

    ins.heading("6.2.6 营销活动、公告提醒与系统设置模块", 3)
    ins.p("营销活动模块主要维护小程序首页展示的横幅内容。管理员可以设置活动标题、图片、排序、状态和有效时间；后端根据状态和时间窗筛选当前可展示的横幅。公告提醒模块用于维护酒店公告，公告可以展示在小程序服务页，并通过开始时间和结束时间控制是否生效。横幅与公告均接入多层时间轮任务，用于处理定时启停。")
    ins.image("activity-manage.png", "图 6-13 营销活动管理界面", 5.8)
    ins.image("notice-manage.png", "图 6-14 公告内容管理界面", 5.8)
    ins.p("系统设置模块维护酒店基础资料、客服电话、服务时间、入住退房规则、地图坐标和封面图片等公共配置。小程序首页、服务页和订单详情页都需要读取这些配置，因此系统设置模块相当于住客端展示内容的后台来源。")

    ins.heading("6.2.7 后台用户、文件上传与日志模块", 3)
    ins.p("后台用户模块用于维护管理员、员工和普通用户的基础信息。管理员可以新增用户、编辑用户状态、修改角色和批量删除用户。文件上传模块主要服务于房型图片、活动横幅、公告图片和用户头像等资源场景；COS 配置完整时可以将图片上传到对象存储，在演示环境中也能保留接口和异常提示。日志模块接收前端错误、接口性能和后台操作线索，后台日志管理页可按照级别、来源和时间查看记录，便于调试和后续运维分析。")
    ins.image("user-manage.png", "图 6-15 后台用户管理界面", 5.8)

    ins.heading("6.3 系统核心业务与支撑能力实现", 2)
    ins.heading("6.3.1 预订、库存与价格联动实现", 3)
    ins.p("预订链路由 BookingService、InventoryService、RoomService 和 UserService 共同完成。创建订单时，BookingService 先校验入住日期、离店日期和房间状态，再调用 InventoryService 检查日期区间内的可售库存；价格计算阶段按天读取库存基础价，并叠加动态定价规则；订单创建后扣减库存。取消订单时释放库存，退房时恢复房态并更新用户积分和累计消费。")

    ins.heading("6.3.2 行为采集与房型推荐实现", 3)
    ins.p("推荐服务通过 user_behaviors 表记录用户行为，将浏览、查看详情、点击推荐和预订意图转换为不同权重的信号。服务层先根据历史行为构建用户与房间的交互矩阵，再结合房间价格、面积、可住人数、折扣率、会员等级和热度等特征进行综合打分，最后返回推荐房型与推荐理由。该实现重点体现行为采集、特征归一化、分数融合和结果解释，不将其夸大为复杂模型训练系统。")

    ins.heading("6.3.3 权限控制、资源上传与日志支撑实现", 3)
    ins.p("权限控制由 JWT 中间件和用户状态共同完成。公共接口允许住客查看房间、公告、横幅和酒店设置；登录后接口允许普通住客查看个人资料和自己的订单；管理员接口则需要通过鉴权后才能访问订单、用户、房间、库存、定价、工单和系统设置等后台能力。资源上传模块封装 COS 上传流程，日志模块封装前端上报和后台查询流程，使核心业务之外的支撑能力也具备清晰边界。")

    ins.heading("6.4 本章小结", 2)
    ins.p("本章按照小程序住客端、管理端运营后台和系统核心支撑能力三个层次，对系统实现进行了逐模块说明。与原稿只把大量功能合并到两个小节不同，本章单独展开了用户认证、首页展示、房型浏览、推荐、在线预订、订单查询、会员服务、运营控制台、前台接待、房间房型、可视化房态、库存矩阵、动态定价、维修清洁、活动公告、系统设置、后台用户、上传日志等模块，使正文结构与实际项目模块更加一致。")


def add_chapter_7(ins):
    ins.page_break()
    ins.heading("第七章  系统测试", 1)
    ins.heading("7.1 测试环境与配置", 2)
    ins.p("系统测试以本地开发环境和演示数据库为基础，重点验证各模块功能是否能够按设计流程运行。测试过程中同时覆盖 Go 后端接口、React 管理端页面和 UniApp 小程序端页面。测试环境如表 7-1 所示。")
    ins.table_caption("表 7-1 测试环境配置")
    ins.table([
        ["测试项", "配置说明"],
        ["操作系统", "Windows 本地开发环境"],
        ["后端环境", "Go、Gin、GORM，使用 MySQL 或 SQLite 演示库"],
        ["管理端环境", "React、Ant Design Pro、浏览器预览"],
        ["小程序端环境", "UniApp H5 预览环境"],
        ["接口测试范围", "用户认证、房间、房型、订单、库存、定价、工单、公告、日志等接口"],
        ["页面测试范围", "小程序端主要页面与管理端运营页面"],
    ])

    ins.heading("7.2 测试方法", 2)
    ins.p("本系统采用功能测试、页面走查、接口验证和业务联动测试相结合的方法。功能测试用于验证单个模块是否可以完成基本操作；页面走查用于检查小程序端和管理端页面是否能正常渲染真实接口数据；接口验证用于检查后端返回结构、分页参数和状态码是否合理；业务联动测试则重点关注订单、库存、价格、房态和会员积分之间的状态变化。")
    ins.p("测试顺序按照“先基础、后模块、再联动”的原则展开。首先验证注册登录和鉴权能力，再分别测试住客端房型浏览、预订、订单、会员和服务页面；随后测试管理端订单、前台接待、房间、房型、库存、定价、工单、公告、用户和日志页面；最后通过完整预订流程检查跨模块状态是否一致。")

    ins.heading("7.3 功能测试", 2)
    ins.heading("7.3.1 住客端功能测试", 3)
    ins.table_caption("表 7-2 住客端功能测试用例")
    ins.table([
        ["测试模块", "测试内容", "预期结果"],
        ["注册登录", "输入有效账号信息注册并登录", "账号创建成功，登录后保存 token 和用户资料。"],
        ["首页展示", "打开首页并加载横幅、酒店信息和推荐房型", "页面展示酒店基础资料、活动横幅和快捷入口。"],
        ["房型浏览", "选择日期后进入房型列表", "展示可售房型、价格、面积、床型和设施信息。"],
        ["房间详情", "进入房间详情并点击推荐房型", "展示房间详情和推荐理由，行为上报接口正常。"],
        ["确认订单", "填写入住人信息并提交订单", "订单创建成功，金额由库存价格和定价规则计算。"],
        ["订单列表", "查看不同状态订单并进入详情", "订单状态、日期、金额和住客信息展示正确。"],
        ["会员中心", "查看个人中心、余额、积分和资料页面", "会员等级、积分、余额和资料维护入口展示正常。"],
        ["酒店服务", "打开服务页查看公告和酒店规则", "展示客服电话、地址、公告、WiFi 和入住退房说明。"],
    ])

    ins.heading("7.3.2 管理端功能测试", 3)
    ins.table_caption("表 7-3 管理端功能测试用例")
    ins.table([
        ["测试模块", "测试内容", "预期结果"],
        ["登录与控制台", "管理员登录后进入运营控制台", "控制台展示入住率、在住间数、待处理工单和快捷入口。"],
        ["订单管理", "查询订单、确认订单、按客人信息搜索", "订单列表与状态操作正常。"],
        ["前台接待", "办理入住和退房", "订单状态、房间状态和清洁状态按流程变化。"],
        ["房间管理", "新增、编辑、删除和批量创建房间", "房号、房型、价格、状态和清洁状态保存正确。"],
        ["房型分类", "新增、修改和删除房型分类", "房型信息保存，引用检查和同步更新有效。"],
        ["房间可视化", "拖拽房间和设施位置", "坐标保存后页面能够按楼层重新渲染。"],
        ["库存矩阵", "初始化库存并按日期区间查询", "库存表格展示剩余房量和基础价格。"],
        ["动态定价", "新增价格规则并查询列表", "规则按房型、日期和优先级保存。"],
        ["维修清洁", "创建维修工单、完成维修、分配清洁任务", "工单状态和房间状态联动正确。"],
        ["活动公告", "创建横幅和公告并设置时间窗", "前台只展示当前有效内容。"],
        ["用户日志设置", "维护后台用户、查看日志、保存酒店设置", "用户状态、日志列表和公共设置保存正常。"],
    ])

    ins.heading("7.3.3 核心业务联动测试", 3)
    ins.table_caption("表 7-4 核心业务联动测试用例")
    ins.table([
        ["测试链路", "操作步骤", "预期结果"],
        ["预订与库存联动", "住客选择日期并创建订单", "订单生成后对应日期库存减少。"],
        ["取消订单联动", "住客取消未入住订单", "订单状态变为 cancelled，对应库存释放。"],
        ["入住退房联动", "后台确认订单、办理入住、办理退房", "订单状态依次变化，房态从可用到入住再恢复可用，清洁状态变为待清洁。"],
        ["动态定价联动", "设置某日期价格规则后创建订单", "订单金额体现规则调整后的每日价格。"],
        ["维修工单联动", "创建维修工单并完成维修", "维修期间房间不可售，完成后恢复可用并进入清洁处理。"],
        ["内容展示联动", "后台创建有效横幅和公告", "小程序首页和服务页能够读取当前有效内容。"],
    ])
    ins.p("小程序端页面走查过程中，首页、房型列表、房间详情、确认订单、订单列表、订单详情、个人中心和酒店服务页面均能加载对应接口数据。管理端页面走查过程中，订单管理、前台接待、房间可视化、库存矩阵、动态定价、工单处理、活动公告、用户管理、日志和系统设置等页面均能进入并展示主要数据。")

    ins.heading("7.4 非功能测试", 2)
    ins.table_caption("表 7-5 非功能测试与质量验证")
    ins.table([
        ["测试类别", "验证内容", "测试结论"],
        ["易用性", "小程序端底部导航、快捷入口、订单状态筛选和后台菜单结构", "主要业务入口清晰，页面层级与功能模块对应。"],
        ["安全性", "JWT 鉴权、TokenVersion、管理员接口保护和账号状态校验", "普通接口和管理接口边界明确，旧令牌可通过版本号失效。"],
        ["一致性", "订单、库存、房态、清洁状态和会员积分联动", "核心状态在典型业务流程中保持一致。"],
        ["可维护性", "handler、service、repository、models 分层结构", "页面、接口、业务规则和数据访问边界清晰。"],
        ["兼容性", "管理端浏览器访问与 UniApp H5 预览", "满足毕业设计演示环境下的基本访问需求。"],
    ])
    ins.image("test-summary-chart.png", "图 7-1 自动化测试结果汇总图", 5.2)
    ins.p("需要说明的是，本文测试主要面向毕业设计演示环境和单酒店典型业务流程，尚未开展真实支付回调、短信通知、多门店部署、大规模并发压测和正式安全渗透测试。因此测试结论主要用于说明系统在当前设计范围内具备基本可用性和模块联动正确性。")

    ins.heading("7.5 本章小结", 2)
    ins.p("本章按照住客端、管理端、核心业务联动和非功能质量四个层面重新组织了测试内容。测试用例覆盖注册登录、首页展示、房型浏览、房间详情、在线预订、订单查询、会员服务、前台接待、房间房型、房态可视化、库存定价、维修清洁、活动公告、用户日志和系统设置等模块，与第四章功能模块和第六章详细实现保持对应。")
    ins.page_break()


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    copy2(SRC, BACKUP)
    doc = Document(str(SRC))
    target = remove_range(doc, "第四章", "第八章")
    ins = Inserter(doc, target)
    add_chapter_4(ins)
    add_chapter_5(ins)
    add_chapter_6(ins)
    add_chapter_7(ins)
    doc.save(str(OUT))
    print(f"saved: {OUT}")
    print(f"backup: {BACKUP}")


if __name__ == "__main__":
    main()
